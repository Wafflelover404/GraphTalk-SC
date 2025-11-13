"""Custom document loaders for handling various file formats."""

import os
import zipfile
import tempfile
import logging
from typing import List, Optional, Dict, Any
from langchain_core.documents import Document
from langchain_community.document_loaders import (
    PyPDFLoader,
    Docx2txtLoader,
    UnstructuredHTMLLoader,
    UnstructuredPDFLoader,
    UnstructuredWordDocumentLoader
)
from langchain.document_loaders.base import BaseLoader

class EnhancedPDFLoader(UnstructuredPDFLoader):
    """Enhanced PDF loader with better text extraction and metadata handling."""
    
    def __init__(self, file_path: str, **kwargs):
        super().__init__(file_path, **kwargs)
        try:
            import pdfplumber  # Better PDF text extraction
            self.use_pdfplumber = True
        except ImportError:
            self.use_pdfplumber = False
            logging.warning("pdfplumber not installed, falling back to basic PDF extraction")
    
    def load(self) -> List[Document]:
        """Load PDF with enhanced extraction if possible."""
        if self.use_pdfplumber:
            try:
                docs = []
                with pdfplumber.open(self.file_path) as pdf:
                    for i, page in enumerate(pdf.pages):
                        text = page.extract_text()
                        if text.strip():  # Skip empty pages
                            metadata = {
                                "source": self.file_path,
                                "page": i + 1,
                                "total_pages": len(pdf.pages),
                                "pdf_width": page.width,
                                "pdf_height": page.height,
                            }
                            docs.append(Document(page_content=text, metadata=metadata))
                return docs
            except Exception as e:
                logging.warning(f"pdfplumber extraction failed, falling back: {str(e)}")
        
        return super().load()

class EnhancedDocxLoader(UnstructuredWordDocumentLoader):
    """Enhanced DOCX loader with better metadata extraction."""
    
    def __init__(self, file_path: str, **kwargs):
        super().__init__(file_path, **kwargs)
        
    def load(self) -> List[Document]:
        """Load DOCX with enhanced metadata extraction."""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(self.file_path)
            
            # Extract document properties
            core_properties = doc.core_properties
            metadata = {
                "source": self.file_path,
                "title": core_properties.title,
                "author": core_properties.author,
                "created": core_properties.created,
                "modified": core_properties.modified,
                "last_modified_by": core_properties.last_modified_by,
                "revision": core_properties.revision,
                "word_count": len(doc.paragraphs)
            }
            
            # Get full text content
            full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            
            return [Document(page_content=full_text, metadata=metadata)]
            
        except ImportError:
            logging.warning("python-docx not installed, falling back to basic DOCX extraction")
            return super().load()

class ZIPLoader(BaseLoader):
    """Loader that handles ZIP files by extracting and loading contained documents."""
    
    def __init__(
        self,
        file_path: str,
        allowed_extensions: Optional[List[str]] = None,
        max_files: int = 100
    ):
        """Initialize the ZIP loader.
        
        Args:
            file_path: Path to the ZIP file
            allowed_extensions: List of allowed file extensions (e.g. ['.pdf', '.docx'])
            max_files: Maximum number of files to process from the ZIP
        """
        self.file_path = file_path
        self.allowed_extensions = allowed_extensions or ['.pdf', '.docx', '.doc', '.txt', '.md', '.html']
        self.max_files = max_files
        self.logger = logging.getLogger(__name__)
    
    def load(self) -> List[Document]:
        """Extract and load documents from the ZIP file."""
        documents = []
        archive_filename = os.path.basename(self.file_path)
        
        self.logger.info(f"Starting ZIP archive processing: {archive_filename}")
        self.logger.info(f"Archive path: {self.file_path}")
        
        extracted_count = 0
        processed_count = 0
        skipped_count = 0
        
        with tempfile.TemporaryDirectory() as temp_dir:
            # Extract ZIP contents
            try:
                with zipfile.ZipFile(self.file_path, 'r') as zip_ref:
                    # Log total files in archive
                    total_files_in_zip = len(zip_ref.filelist)
                    self.logger.info(f"Total files in archive: {total_files_in_zip}")
                    
                    # Filter files by extension and limit
                    files_to_extract = []
                    for file_info in zip_ref.filelist:
                        if len(files_to_extract) >= self.max_files:
                            self.logger.warning(f"Reached max file limit of {self.max_files}, stopping extraction")
                            break
                        if any(file_info.filename.lower().endswith(ext) for ext in self.allowed_extensions):
                            files_to_extract.append(file_info.filename)
                    
                    self.logger.info(f"Files to extract: {len(files_to_extract)} (Allowed extensions: {', '.join(self.allowed_extensions)})")
                    
                    # Extract selected files
                    zip_ref.extractall(temp_dir, members=files_to_extract)
                    extracted_count = len(files_to_extract)
                    self.logger.info(f"Successfully extracted {extracted_count} files to temporary directory")
                    
                    # Log extracted files
                    for file_name in files_to_extract:
                        self.logger.debug(f"Extracted: {file_name}")
            
            except zipfile.BadZipFile as e:
                self.logger.error(f"Invalid ZIP file: {archive_filename} - {str(e)}")
                raise
            except Exception as e:
                self.logger.error(f"Error extracting ZIP archive {archive_filename}: {str(e)}")
                raise
            
            # Process each extracted file
            for root, _, files in os.walk(temp_dir):
                for filename in files:
                    file_path = os.path.join(root, filename)
                    ext = os.path.splitext(filename)[1].lower()
                    
                    try:
                        self.logger.info(f"Processing extracted file: {filename} (Type: {ext})")
                        
                        if ext == '.pdf':
                            loader = EnhancedPDFLoader(file_path)
                        elif ext in ['.docx', '.doc']:
                            loader = EnhancedDocxLoader(file_path)
                        elif ext == '.html':
                            loader = UnstructuredHTMLLoader(file_path)
                        elif ext in ['.txt', '.md']:
                            with open(file_path, 'r', encoding='utf-8') as f:
                                docs = [Document(
                                    page_content=f.read(),
                                    metadata={"source": f"zip://{self.file_path}/{filename}"}
                                )]
                        else:
                            self.logger.debug(f"Skipping unsupported file type: {filename}")
                            skipped_count += 1
                            continue
                        
                        if not 'docs' in locals():
                            docs = loader.load()
                        
                        # Update metadata to indicate ZIP source
                        for doc in docs:
                            doc.metadata["archive_source"] = self.file_path
                            doc.metadata["archive_filename"] = archive_filename
                            doc.metadata["archive_path"] = filename
                        
                        chunk_count = len(docs)
                        documents.extend(docs)
                        processed_count += 1
                        
                        self.logger.info(f"✓ Processed {filename}: {chunk_count} document(s) loaded from archive")
                        
                    except Exception as e:
                        self.logger.error(f"Error processing {filename} from ZIP {archive_filename}: {str(e)}")
                        continue
        
        # Log summary
        self.logger.info(f"ZIP PROCESSING SUMMARY for {archive_filename}")
        self.logger.info(f"  - Extracted files: {extracted_count}")
        self.logger.info(f"  - Successfully processed: {processed_count}")
        self.logger.info(f"  - Skipped files: {skipped_count}")
        self.logger.info(f"  - Total documents loaded: {len(documents)}")
        self.logger.info(f"Archive {archive_filename} processing complete")
        
        return documents

# Example usage:
# zip_loader = ZIPLoader("documents.zip")
# documents = zip_loader.load()